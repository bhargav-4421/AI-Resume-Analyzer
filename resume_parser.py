import pdfplumber
import docx
import os

def extract_resume_text(file):
    """
    Extract text from PDF and DOCX resume files.
    
    Args:
        file: File path or file-like object
        
    Returns:
        str: Extracted text from the resume
        
    Raises:
        ValueError: If file type is not supported
        Exception: If file extraction fails
    """
    try:
        # Handle file path string
        if isinstance(file, str):
            if not os.path.exists(file):
                raise FileNotFoundError(f"File not found: {file}")
            
            file_extension = os.path.splitext(file)[1].lower()
            
            if file_extension == '.pdf':
                return extract_from_pdf(file)
            elif file_extension in ['.docx', '.doc']:
                return extract_from_docx(file)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Handle file-like object (e.g., from FileStorage in Flask)
        else:
            filename = getattr(file, 'filename', '')
            if not filename:
                raise ValueError("No filename provided for file-like object")
            
            file_extension = os.path.splitext(filename)[1].lower()
            
            if file_extension == '.pdf':
                return extract_from_pdf_fileobj(file)
            elif file_extension in ['.docx', '.doc']:
                return extract_from_docx_fileobj(file)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
    except Exception as e:
        raise Exception(f"Error extracting text from resume: {str(e)}")

def extract_from_pdf(file_path):
    """Extract text from PDF file using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Error reading PDF file: {str(e)}")

def extract_from_pdf_fileobj(fileobj):
    """Extract text from PDF file-like object using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(fileobj) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Error reading PDF file object: {str(e)}")

def extract_from_docx(file_path):
    """Extract text from DOCX file using python-docx."""
    try:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {str(e)}")

def extract_from_docx_fileobj(fileobj):
    """Extract text from DOCX file-like object using python-docx."""
    try:
        doc = docx.Document(fileobj)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Error reading DOCX file object: {str(e)}")

# Example usage and testing
if __name__ == "__main__":
    # Test with file paths
    # pdf_text = extract_resume_text("resume.pdf")
    # docx_text = extract_resume_text("resume.docx")
    # print("PDF Text:", pdf_text[:100] + "...")
    # print("DOCX Text:", docx_text[:100] + "...")
    pass